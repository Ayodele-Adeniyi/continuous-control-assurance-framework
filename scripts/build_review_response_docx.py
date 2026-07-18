"""Build the concise Word review-response form from REVIEW_RESPONSE_TEMPLATE.md."""

from __future__ import annotations

import re
from argparse import ArgumentParser
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "REVIEW_RESPONSE_TEMPLATE.md"
OUTPUT = ROOT / "docs" / "CCAF_Independent_Review_Response_Template.docx"

NAVY = "193A5A"
LIGHT = "EAF0F5"
MID = "8EA2B3"
WHITE = "FFFFFF"
TEXT = RGBColor(28, 39, 49)
MUTED = RGBColor(86, 98, 108)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(
    cell, top: int = 65, start: int = 90, bottom: int = 65, end: int = 90
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = MID, size: str = "5") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_table_width(table, widths: tuple[float, ...]) -> None:
    table.autofit = False
    for column, width in zip(table.columns, widths):
        column.width = Inches(width)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(col)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(sum(widths) * 1440)))
    tbl_w.set(qn("w:type"), "dxa")


def add_text(
    paragraph,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    size: float = 8.2,
    color=TEXT,
):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return run


def add_section_heading(doc: Document, number: str, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(2.5)
    paragraph.paragraph_format.space_after = Pt(1.5)
    paragraph.paragraph_format.keep_with_next = True
    add_text(paragraph, f"{number}. {text}", bold=True, size=10.2, color=RGBColor(25, 58, 90))


def add_opinion_group(
    doc: Document,
    title: str,
    lead: str,
    options: tuple[str, str, str],
) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(1.2)
    paragraph.paragraph_format.space_after = Pt(0.5)
    paragraph.paragraph_format.keep_with_next = True
    add_text(paragraph, f"{title}. ", bold=True, size=8.4)
    add_text(paragraph, lead, size=8.1)

    choices = doc.add_paragraph()
    choices.paragraph_format.left_indent = Inches(0.08)
    choices.paragraph_format.space_after = Pt(0.3)
    choices.paragraph_format.line_spacing = 0.92
    for index, option in enumerate(options):
        if index:
            choices.add_run().add_break()
        add_text(choices, f"☐ {option}", size=7.9)

    comments = doc.add_paragraph()
    comments.paragraph_format.space_after = Pt(0.7)
    add_text(comments, "Comments (optional): ", italic=True, size=7.8, color=MUTED)
    add_text(comments, "_______________________________________________________________", size=7.8, color=MUTED)


def source_values() -> tuple[str, str, str]:
    text = SOURCE.read_text(encoding="utf-8")
    snapshot = re.search(r"\*\*Stable review snapshot:\*\* `([^`]+)`", text)
    url = re.search(r"\*\*Snapshot URL:\*\* (https://\S+)", text)
    facts = re.search(r"\*\*Documented release facts:\*\* (.+)", text)
    if not snapshot or not url or not facts:
        raise ValueError("Required snapshot and release facts are missing from the Markdown source")
    return snapshot.group(1), url.group(1), facts.group(1)


def build(output: Path = OUTPUT) -> Path:
    snapshot, snapshot_url, release_facts = source_values()
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.38)
    section.bottom_margin = Inches(0.38)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(8.2)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_after = Pt(1)
    normal.paragraph_format.line_spacing = 1.0

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(0.5)
    add_text(
        title,
        "INDEPENDENT TECHNICAL REVIEW RESPONSE",
        bold=True,
        size=13.3,
        color=RGBColor(25, 58, 90),
    )
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(3)
    add_text(
        subtitle,
        "Continuous Control Assurance Framework (CCAF), Version 1.3.1",
        italic=True,
        size=8.8,
    )

    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(2)
    add_text(
        intro,
        "This concise form records the materials and procedures personally reviewed and the reviewer's independent professional judgment. A documents-based design review is sufficient; repository inspection and local reproduction are optional.",
        size=7.8,
        color=MUTED,
    )

    add_section_heading(doc, "1", "Reviewer information")
    reviewer = doc.add_table(rows=4, cols=2)
    set_table_width(reviewer, (3.70, 3.70))
    set_table_borders(reviewer)
    reviewer_rows = (
        ("Reviewer name: __________________________", "Current professional title: __________________________"),
        ("Organization: __________________________", "Email or telephone: __________________________"),
        ("Relevant experience (brief): __________________", "Review date: __________________________"),
    )
    for row, values in zip(reviewer.rows[:3], reviewer_rows):
        for cell, value in zip(row.cells, values):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            add_text(paragraph, value, bold=True, size=7.9)
    conflict = reviewer.rows[3].cells[0].merge(reviewer.rows[3].cells[1])
    set_cell_margins(conflict)
    add_text(
        conflict.paragraphs[0],
        "Prior relationship, conflict, or compensation, if any: ________________________________________________",
        bold=True,
        size=7.9,
    )

    add_section_heading(doc, "2", "Materials, scope, and procedures reviewed")
    record = doc.add_table(rows=4, cols=1)
    set_table_width(record, (7.40,))
    set_table_borders(record)

    p = record.rows[0].cells[0].paragraphs[0]
    set_cell_shading(record.rows[0].cells[0], LIGHT)
    set_cell_margins(record.rows[0].cells[0], top=55, bottom=55)
    add_text(p, f"Stable snapshot: {snapshot}  |  ", bold=True, size=7.7)
    add_text(p, snapshot_url, size=7.25, color=RGBColor(33, 91, 145))

    p = record.rows[1].cells[0].paragraphs[0]
    set_cell_margins(record.rows[1].cells[0], top=45, bottom=45)
    add_text(p, "Materials/procedures: ", bold=True, size=7.7)
    add_text(
        p,
        "☐ Methodology  ☐ Control-Test Catalog  ☐ Repository/code (optional)  ☐ Local reproduction (optional)  ☐ Automated tests (optional)",
        size=7.5,
    )

    p = record.rows[2].cells[0].paragraphs[0]
    set_cell_margins(record.rows[2].cells[0], top=45, bottom=45)
    add_text(p, "Catalog scope: ", bold=True, size=7.7)
    add_text(p, "☐ All 20 controls  ☐ Selected modules/control IDs: __________________________", size=7.6)
    p.add_run().add_break()
    add_text(p, "Observed result and seed, if locally reproduced: __________________________________________", size=7.5)

    p = record.rows[3].cells[0].paragraphs[0]
    set_cell_shading(record.rows[3].cells[0], LIGHT)
    set_cell_margins(record.rows[3].cells[0], top=55, bottom=55)
    add_text(p, "Documented release facts: ", bold=True, size=7.55)
    add_text(p, release_facts, size=7.4)
    p.add_run().add_break()
    add_text(p, "Corrections, if any: ________________________________________________", italic=True, size=7.4, color=MUTED)

    add_section_heading(doc, "3", "Professional opinion")
    add_opinion_group(
        doc,
        "Technical soundness",
        "In my professional judgment, this framework is:",
        (
            "Technically sound and consistent with professional control-testing practice for a synthetic reference prototype",
            "Technically sound with minor reservations explained below",
            "Not technically sound, as explained below",
        ),
    )
    add_opinion_group(
        doc,
        "Claims and limitations",
        "In my professional judgment, the framework's claims are:",
        (
            "Appropriately bounded to what the demonstration establishes",
            "Generally appropriate, but some claims are overstated as explained below",
            "Not appropriately bounded, as explained below",
        ),
    )
    add_opinion_group(
        doc,
        "Transferability",
        "In my professional judgment, this methodology:",
        (
            "Could be adapted by other institutions or practitioners to their own authorized data and control environments",
            "Could be adapted with significant effort, as explained below",
            "Is not practically transferable, as explained below",
        ),
    )

    overall = doc.add_paragraph()
    overall.paragraph_format.space_before = Pt(1)
    overall.paragraph_format.space_after = Pt(0.5)
    overall.paragraph_format.keep_with_next = True
    add_text(overall, "Overall assessment. ", bold=True, size=8.4)
    add_text(
        overall,
        "Briefly address technical soundness, limitations, and potential transferability (2-4 sentences):",
        size=8.0,
    )
    response_box = doc.add_table(rows=1, cols=1)
    set_table_width(response_box, (7.40,))
    set_table_borders(response_box)
    response_box.rows[0].height = Inches(0.58)
    response_box.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    response_cell = response_box.rows[0].cells[0]
    response_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_margins(response_cell, top=55, bottom=55)
    add_text(response_cell.paragraphs[0], "", size=8.0)

    add_section_heading(doc, "4", "Confirmation")
    confirmation = doc.add_paragraph()
    confirmation.paragraph_format.space_after = Pt(1)
    add_text(
        confirmation,
        "This response reflects my independent professional judgment based on the materials and procedures identified above. It does not constitute institutional endorsement, regulatory approval, certification, a claim of institutional adoption, or a claim about production performance. I understand that this response may be cited publicly and in professional or immigration-related submissions as evidence of independent review.",
        size=7.45,
    )

    signature = doc.add_table(rows=1, cols=3)
    set_table_width(signature, (2.25, 3.15, 2.00))
    for cell, value in zip(
        signature.rows[0].cells,
        (
            "Name (printed): __________________",
            "Signature: __________________________",
            "Date: ______________",
        ),
    ):
        set_cell_margins(cell, top=55, bottom=55)
        add_text(cell.paragraphs[0], value, bold=True, size=8.0)
    set_table_borders(signature, color=WHITE, size="0")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(footer, f"CCAF independent review | Stable snapshot: {snapshot}", size=7.0, color=MUTED)

    doc.core_properties.title = "CCAF Independent Technical Review Response"
    doc.core_properties.subject = "Concise independent technical review response form"
    doc.core_properties.author = "Ayodele Timothy Adeniyi"
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    return output


if __name__ == "__main__":
    parser = ArgumentParser(description=__doc__)
    parser.add_argument("--output", type=Path, default=OUTPUT)
    args = parser.parse_args()
    print(build(args.output))
